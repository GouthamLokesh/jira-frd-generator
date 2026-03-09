import streamlit as st
import os
import requests
from requests.auth import HTTPBasicAuth
from dotenv import load_dotenv
import docx
from openai import OpenAI
import google.generativeai as genai
import json
import io
import PyPDF2

# Load environment variables from the user's specific file
load_dotenv("APIKey.env")

# Configuration
JIRA_BASE_URL = os.getenv("JIRA_BASE_URL")
JIRA_EMAIL = os.getenv("JIRA_EMAIL")
JIRA_API_TOKEN = os.getenv("JIRA_API_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

st.set_page_config(page_title="Jira to FRD Generator", page_icon="📝", layout="wide")

# Premium animated CSS
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

    /* --- Keyframe Animations --- */
    @keyframes gradientShift {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }
    @keyframes float {
        0%, 100% { transform: translateY(0px) rotate(0deg); }
        33% { transform: translateY(-18px) rotate(2deg); }
        66% { transform: translateY(-8px) rotate(-1deg); }
    }
    @keyframes pulse-glow {
        0%, 100% { box-shadow: 0 0 20px rgba(139, 92, 246, 0.3), 0 0 40px rgba(59, 130, 246, 0.1); }
        50% { box-shadow: 0 0 40px rgba(139, 92, 246, 0.6), 0 0 80px rgba(59, 130, 246, 0.3); }
    }
    @keyframes shimmer {
        0% { background-position: -200% center; }
        100% { background-position: 200% center; }
    }
    @keyframes fadeInUp {
        from { opacity: 0; transform: translateY(30px); }
        to { opacity: 1; transform: translateY(0); }
    }
    @keyframes textGlow {
        0%, 100% { text-shadow: 0 0 10px rgba(168, 85, 247, 0.4), 0 0 20px rgba(59, 130, 246, 0.2); }
        50% { text-shadow: 0 0 20px rgba(168, 85, 247, 0.8), 0 0 40px rgba(59, 130, 246, 0.5), 0 0 60px rgba(168, 85, 247, 0.3); }
    }
    @keyframes spin-slow {
        from { transform: rotate(0deg); }
        to { transform: rotate(360deg); }
    }
    @keyframes borderPulse {
        0%, 100% { border-color: rgba(139, 92, 246, 0.3); }
        50% { border-color: rgba(139, 92, 246, 0.8); }
    }
    @keyframes dotBounce {
        0%, 80%, 100% { transform: scale(0.6); opacity: 0.4; }
        40% { transform: scale(1.0); opacity: 1; }
    }
    @keyframes orb1 {
        0%, 100% { transform: translate(0, 0) scale(1); }
        50% { transform: translate(40px, -30px) scale(1.15); }
    }
    @keyframes orb2 {
        0%, 100% { transform: translate(0, 0) scale(1); }
        50% { transform: translate(-30px, 40px) scale(0.9); }
    }

    /* --- Base Styles --- */
    body, .stApp {
        background: #020817 !important;
        color: #f1f5f9 !important;
        font-family: 'Inter', sans-serif !important;
    }

    /* Animated Background Orbs */
    .stApp::before {
        content: '';
        position: fixed;
        top: -150px;
        left: -150px;
        width: 500px;
        height: 500px;
        border-radius: 50%;
        background: radial-gradient(circle, rgba(139, 92, 246, 0.15), transparent 70%);
        animation: orb1 10s ease-in-out infinite;
        pointer-events: none;
        z-index: 0;
    }
    .stApp::after {
        content: '';
        position: fixed;
        bottom: -100px;
        right: -100px;
        width: 400px;
        height: 400px;
        border-radius: 50%;
        background: radial-gradient(circle, rgba(59, 130, 246, 0.12), transparent 70%);
        animation: orb2 12s ease-in-out infinite;
        pointer-events: none;
        z-index: 0;
    }

    /* --- Title Block --- */
    .title-wrapper {
        text-align: center;
        padding: 3rem 1rem 2rem;
        animation: fadeInUp 0.8s ease forwards;
    }
    .badge {
        display: inline-block;
        padding: 0.3rem 1rem;
        border-radius: 9999px;
        background: linear-gradient(135deg, rgba(139,92,246,0.2), rgba(59,130,246,0.2));
        border: 1px solid rgba(139,92,246,0.4);
        color: #c4b5fd;
        font-size: 0.8rem;
        font-weight: 600;
        letter-spacing: 0.1em;
        text-transform: uppercase;
        margin-bottom: 1rem;
    }
    .main-title {
        font-size: 3.8rem !important;
        font-weight: 800 !important;
        letter-spacing: -0.03em;
        background: linear-gradient(135deg, #e879f9 0%, #a855f7 25%, #3b82f6 75%, #06b6d4 100%);
        background-size: 200% auto;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        animation: shimmer 4s linear infinite, textGlow 3s ease-in-out infinite;
        margin-bottom: 0.75rem !important;
        line-height: 1.1 !important;
    }
    .subtitle {
        color: #94a3b8;
        font-size: 1.1rem;
        font-weight: 400;
        max-width: 600px;
        margin: 0 auto 0.5rem;
        line-height: 1.6;
    }
    .divider {
        width: 80px;
        height: 3px;
        background: linear-gradient(90deg, #a855f7, #3b82f6);
        margin: 1.5rem auto;
        border-radius: 9999px;
    }

    /* --- Form Card --- */
    [data-testid="stForm"] {
        background: linear-gradient(145deg, rgba(15,23,42,0.9), rgba(30,41,59,0.8)) !important;
        border-radius: 24px !important;
        border: 1px solid rgba(139, 92, 246, 0.25) !important;
        padding: 2.5rem !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        box-shadow: 
            0 25px 50px -12px rgba(0, 0, 0, 0.6),
            0 0 0 1px rgba(139, 92, 246, 0.1),
            inset 0 1px 0 rgba(255,255,255,0.05) !important;
        animation: pulse-glow 4s ease-in-out infinite, fadeInUp 0.9s ease 0.2s both !important;
        transition: border-color 0.3s ease, transform 0.3s ease !important;
        position: relative;
        overflow: hidden;
    }
    [data-testid="stForm"]::before {
        content: '';
        position: absolute;
        top: 0; left: 0; right: 0; bottom: 0;
        background: linear-gradient(135deg, rgba(139,92,246,0.03) 0%, transparent 50%, rgba(59,130,246,0.03) 100%);
        pointer-events: none;
    }
    [data-testid="stForm"]:hover {
        transform: translateY(-4px) !important;
        border-color: rgba(139, 92, 246, 0.6) !important;
    }

    /* --- Inputs --- */
    .stTextInput > div > div > input {
        background: rgba(2, 8, 23, 0.8) !important;
        border: 1px solid rgba(100, 116, 139, 0.4) !important;
        color: #f1f5f9 !important;
        border-radius: 12px !important;
        padding: 0.875rem 1.25rem !important;
        font-size: 1rem !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
    }
    .stTextInput > div > div > input:focus {
        border-color: #8b5cf6 !important;
        box-shadow: 0 0 0 3px rgba(139, 92, 246, 0.2), 0 0 20px rgba(139, 92, 246, 0.15) !important;
        background: rgba(15, 10, 30, 0.9) !important;
    }
    .stTextInput > label {
        color: #94a3b8 !important;
        font-weight: 600 !important;
        font-size: 0.9rem !important;
        letter-spacing: 0.05em;
        text-transform: uppercase;
    }

    /* --- Generate Button --- */
    .stButton > button {
        width: 100%;
        background: linear-gradient(135deg, #8b5cf6 0%, #6366f1 50%, #3b82f6 100%) !important;
        background-size: 200% auto !important;
        color: white !important;
        font-weight: 700 !important;
        font-size: 1rem !important;
        border-radius: 12px !important;
        border: none !important;
        padding: 0.875rem 1.5rem !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        letter-spacing: 0.025em;
        position: relative;
        overflow: hidden;
        animation: shimmer 3s linear infinite !important;
    }
    .stButton > button:hover {
        transform: translateY(-2px) scale(1.01) !important;
        box-shadow: 0 10px 30px rgba(139, 92, 246, 0.5), 0 0 40px rgba(99, 102, 241, 0.3) !important;
        background-position: right center !important;
    }
    .stButton > button:active {
        transform: translateY(0) scale(0.99) !important;
    }

    /* --- Download Button --- */
    .stDownloadButton > button {
        width: 100%;
        background: linear-gradient(135deg, #10b981 0%, #059669 50%, #047857 100%) !important;
        color: white !important;
        font-weight: 700 !important;
        font-size: 1rem !important;
        border-radius: 12px !important;
        border: none !important;
        padding: 0.875rem 1.5rem !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        margin-top: 1.25rem;
        animation: fadeInUp 0.5s ease forwards;
    }
    .stDownloadButton > button:hover {
        transform: translateY(-2px) scale(1.01) !important;
        box-shadow: 0 10px 30px rgba(16, 185, 129, 0.5), 0 0 40px rgba(5, 150, 105, 0.3) !important;
    }

    /* --- Status Messages --- */
    [data-testid="stSuccess"], .stSuccess {
        background: linear-gradient(135deg, rgba(16,185,129,0.1), rgba(5,150,105,0.05)) !important;
        border: 1px solid rgba(16, 185, 129, 0.3) !important;
        border-radius: 12px !important;
        color: #34d399 !important;
        animation: fadeInUp 0.4s ease !important;
    }
    [data-testid="stAlert"], .stError {
        border-radius: 12px !important;
        animation: fadeInUp 0.4s ease !important;
    }
    [data-testid="stWarning"], .stWarning {
        border-radius: 12px !important;
        animation: fadeInUp 0.4s ease !important;
    }

    /* --- Spinner --- */
    .stSpinner > div > div {
        border-top-color: #8b5cf6 !important;
        animation: spin-slow 0.8s linear infinite !important;
    }
    
    /* --- Scrollbar --- */
    ::-webkit-scrollbar { width: 6px; }
    ::-webkit-scrollbar-track { background: #0f172a; }
    ::-webkit-scrollbar-thumb { background: rgba(139,92,246,0.5); border-radius: 9999px; }
    ::-webkit-scrollbar-thumb:hover { background: rgba(139,92,246,0.8); }

    /* --- Headers --- */
    h1, h2, h3 { color: #ffffff !important; font-weight: 700 !important; }
</style>
""", unsafe_allow_html=True)

# Animated Title HTML
st.markdown("""
<div class="title-wrapper">
    <div class="badge">✦ AI-Powered Document Generator</div>
    <h1 class="main-title">FRD Generator</h1>
    <p class="subtitle">Enter a Jira Case ID and let AI automatically build a complete Functional Requirements Document from your case details, description, and attachments.</p>
    <div class="divider"></div>
</div>
""", unsafe_allow_html=True)


def get_jira_issue(issue_id):
    if not all([JIRA_BASE_URL, JIRA_EMAIL, JIRA_API_TOKEN]):
        st.error("Jira credentials not fully configured in .env file.")
        return None
    
    url = f"{JIRA_BASE_URL.rstrip('/')}/rest/api/2/issue/{issue_id}"
    auth = HTTPBasicAuth(JIRA_EMAIL, JIRA_API_TOKEN)
    headers = {"Accept": "application/json"}
    
    response = requests.get(url, headers=headers, auth=auth)
    
    if response.status_code == 200:
        return response.json()
    else:
        st.error(f"Failed to fetch Jira ID {issue_id}. Status Code: {response.status_code}. Response: {response.text}")
        return None

def download_and_parse_attachment(url, filename, auth):
    response = requests.get(url, auth=auth)
    if response.status_code != 200:
        return ""
    
    content = ""
    ext = filename.split('.')[-1].lower()
    
    try:
        if ext == 'txt':
            content = response.text
        elif ext == 'pdf':
            pdf_file = io.BytesIO(response.content)
            reader = PyPDF2.PdfReader(pdf_file)
            for page in reader.pages:
                content += page.extract_text() + "\n"
        elif ext == 'docx':
            doc_file = io.BytesIO(response.content)
            doc = docx.Document(doc_file)
            content = "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        content = f"[Failed to parse attachment {filename}: {str(e)}]"
        
    return f"--- Attachment: {filename} ---\n{content}\n--- End of Attachment ---\n"

def generate_frd_mapping(summary, description, attachments_text):
    if not OPENAI_API_KEY and not GEMINI_API_KEY:
        st.error("Neither OpenAI API Key nor Gemini API Key is configured in the .env file.")
        return None
        
    prompt = f"""
    You are an expert Business Analyst writing a Functional Requirements Document (FRD).
    Based on the following Jira case details and attachments, generate the required text for an FRD template.

    Jira Summary: {summary}
    Jira Description: {description}
    Attachments Content:
    {attachments_text}

    We need you to output a JSON object mapping specific placeholder keys to the text you want to substitute into the Word document.
    Generate a JSON object with the exact following keys and populate them appropriately based on the Jira context. If context is missing for a specific field, make a reasonable assumption or state that it is not specified in the Jira ticket.
    
    Required JSON keys:
    "[Provide introduction/summary of the requirement]"
    "[Insert Feature Toggle documentation to explain whether a toggle in configuration is required for the feature]"
    "[Insert Developer Notes that span multiple Use Cases]"
    "[Use Case Name]"
    "[As a <insert operator role/persona>, I want <insert necessary action/demand>, so that <insert desired value/outcome>]"
    "[Insert bulleted list of assumptions]"
    "[Insert bulleted list of preconditions]"
    "[Should be in a user persona, contain a Beginning, Middle, & End should always leave User at the Beginning]"
    "[Step 1]"
    "[Step 2]"
    "[Step 2.1]"
    "[Step 2.2]"
    "[Insert bulleted list of business rules]"
    "[Insert bulleted list of developer notes]"
    "[Insert process flow]"
    "[Insert wireframe or link to wireframes]"
    "Screen Components": [
       {{ "Label": "...", "Type": "...", "Default": "...", "Enabled": "True/False", "Required": "True/False", "Notes": "..." }}
    ]
    
    Respond ONLY with the JSON format, without markdown code blocks, to ensure it can be parsed with json.loads(). Output standard strings (use \\n for newlines within strings if needed). For the `Screen Components` key, return a list of JSON objects matching the schema above.
    """

    # Try OpenAI first if available
    if OPENAI_API_KEY:
        try:
            client = OpenAI(api_key=OPENAI_API_KEY)
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that outputs JSON."},
                    {"role": "user", "content": prompt}
                ],
                response_format={ "type": "json_object" }
            )
            content = response.choices[0].message.content
            return json.loads(content)
        except Exception as e:
            if "insufficient_quota" in str(e).lower() or "429" in str(e):
                if GEMINI_API_KEY:
                    st.warning("OpenAI quota exceeded. Falling back to Gemini...")
                else:
                    st.error("OpenAI quota exceeded and no Gemini API Key provided for fallback.")
                    return None
            else:
                st.error(f"Error calling OpenAI API: {str(e)}")
                return None

    # Fallback to Gemini if OpenAI failed or wasn't provided
    if GEMINI_API_KEY:
        try:
            genai.configure(api_key=GEMINI_API_KEY)
            model = genai.GenerativeModel('gemini-2.5-flash', generation_config={"response_mime_type": "application/json"})
            response = model.generate_content(prompt)
            return json.loads(response.text)
        except Exception as e:
            st.error(f"Error calling Gemini API: {str(e)}")
            return None
            
    return None

def process_document(template_path, mapping, jira_url, case_summary):
    doc = docx.Document(template_path)
    
    # Add standard mappings
    mapping["[Insert Jira ticket/link]"] = jira_url
    mapping["Title"] = case_summary
    
    # --- STEP 1: Process tables FIRST ---
    # Handle Screen Components table specially, then do text replacement for all others
    for table in doc.tables:
        # Check if this is the Screen Components table (6-column table with specific header)
        if len(table.rows) > 0 and len(table.rows[0].cells) == 6:
            header_text = [cell.text.strip() for cell in table.rows[0].cells]
            if header_text == ['Label', 'Type', 'Default', 'Enabled', 'Required', 'Notes']:
                # Populate with AI-generated component data
                components = mapping.get("Screen Components", [])
                if isinstance(components, list) and components:
                    # Remove the template placeholder row (row 1)
                    if len(table.rows) > 1:
                        table._tbl.remove(table.rows[1]._tr)
                    # Add a new row for each component
                    for comp in components:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(comp.get("Label", ""))
                        row_cells[1].text = str(comp.get("Type", ""))
                        row_cells[2].text = str(comp.get("Default", ""))
                        row_cells[3].text = str(comp.get("Enabled", ""))
                        row_cells[4].text = str(comp.get("Required", ""))
                        row_cells[5].text = str(comp.get("Notes", ""))
                        
                # Remove from mapping so paragraphs don't get the JSON text dumped
                mapping.pop("Screen Components", None)
                continue  # skip normal text replacement for this table

        # Normal text replacement for all other tables
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in mapping.items():
                        if key in paragraph.text:
                            if key == paragraph.text.strip():
                                paragraph.text = str(value)
                            else:
                                paragraph.text = paragraph.text.replace(key, str(value))

    # --- STEP 2: Process paragraphs AFTER tables ---
    # "Screen Components" key has already been removed, so heading text won't be replaced
    for paragraph in doc.paragraphs:
        for key, value in mapping.items():
            if key in paragraph.text:
                if key == paragraph.text.strip():
                    paragraph.text = str(value)
                else:
                    paragraph.text = paragraph.text.replace(key, str(value))

    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

col1, col2, col3 = st.columns([1, 2, 1])

with col2:
    with st.form("jira_form"):
        jira_id = st.text_input("📦 Jira Case ID (e.g., PROJ-123)", placeholder="Enter ID...")
        submit_button = st.form_submit_button("Generate FRD ✨")

    if submit_button:
        if not jira_id.strip():
            st.warning("Please enter a valid Jira Case ID.")
        else:
            with st.spinner("Fetching Jira Case..."):
                issue_data = get_jira_issue(jira_id.strip())
                
            if issue_data:
                summary = issue_data.get("fields", {}).get("summary", "")
                description = issue_data.get("fields", {}).get("description", "")
                
                # Handle varying description formats
                if isinstance(description, dict):
                    description = json.dumps(description)
                elif not description:
                    description = "No description provided."
                
                st.success(f"Fetched Case: {summary}")
                
                attachments = issue_data.get("fields", {}).get("attachment", [])
                attachments_text = ""
                
                if attachments:
                    with st.spinner(f"Processing {len(attachments)} attachment(s)..."):
                        auth = HTTPBasicAuth(JIRA_EMAIL, JIRA_API_TOKEN)
                        for att in attachments:
                            url = att.get("content")
                            filename = att.get("filename", "")
                            if any(filename.lower().endswith(ext) for ext in ['.txt', '.pdf', '.docx']):
                                attachments_text += download_and_parse_attachment(url, filename, auth)
                
                with st.spinner("Analyzing and Generating FRD content with OpenAI..."):
                    mapping = generate_frd_mapping(summary, description, attachments_text)
                    
                if mapping:
                    with st.spinner("Populating DOCX Template..."):
                        template_path = "FRD template.docx"
                        if not os.path.exists(template_path):
                            st.error(f"Template file '{template_path}' not found in the directory.")
                        else:
                            try:
                                jira_link = f"{JIRA_BASE_URL.rstrip('/')}/browse/{jira_id.strip()}" if JIRA_BASE_URL else jira_id.strip()
                                output_stream = process_document(template_path, mapping, jira_link, summary)
                                
                                st.success("🎉 FRD Generated Successfully!")
                                st.download_button(
                                    label="Download Generated FRD (.docx) ⬇️",
                                    data=output_stream,
                                    file_name=f"FRD_{jira_id.strip()}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                            except Exception as e:
                                st.error(f"Error processing document: {str(e)}")
